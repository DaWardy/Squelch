"""Generate Squelch bootstrap DOCX documents."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy
from pathlib import Path

OUT = Path(__file__).parent

# ── Helpers ──────────────────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def _set_cell_border(cell, color="CCCCCC"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)

def add_heading(doc, text, level=1, color=None):
    h = doc.add_heading(text, level=level)
    if color:
        for run in h.runs:
            run.font.color.rgb = RGBColor(*color)
    return h

def add_box(doc, text, bg_hex, border_hex, label=None):
    """Add a colored notice box paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.right_indent = Inches(0.2)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    # shade via XML
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), bg_hex)
    pPr.append(shd)
    if label:
        run = p.add_run(f"{label}  ")
        run.bold = True
        run.font.color.rgb = RGBColor(*bytes.fromhex(border_hex))
    run2 = p.add_run(text)
    run2.font.size = Pt(10)
    return p

def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F5F5F5')
    pPr.append(shd)
    return p

def add_table(doc, headers, rows, col_widths=None, header_bg="2E5FA3"):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _set_cell_bg(hdr_cells[i], header_bg)
        if col_widths:
            hdr_cells[i].width = Inches(col_widths[i])
    # Data rows
    for ri, row in enumerate(rows):
        row_cells = table.add_row().cells
        bg = "EFF3FB" if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row):
            row_cells[ci].text = str(val)
            row_cells[ci].paragraphs[0].runs[0].font.size = Pt(9)
            _set_cell_bg(row_cells[ci], bg)
            _set_cell_border(row_cells[ci])
            if col_widths:
                row_cells[ci].width = Inches(col_widths[ci])
    return table

def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def add_numbered(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Number')
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

# ── Document 1: Bootstrap Template (INTERNAL) ────────────────────────────────

def build_bootstrap():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.25)
        section.right_margin  = Inches(1.25)

    # Title
    title = doc.add_heading('Squelch — Project Bootstrap Template', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph('AI Agent and Developer Reference  |  INTERNAL ONLY')
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    sub.runs[0].font.size = Pt(12)
    doc.add_paragraph()

    # RED warning box
    add_box(doc,
        "INTERNAL USE ONLY — Do not commit to the public repository. "
        "Contains stakeholder personas, operational context, and threat model.",
        "FFE0E0", "CC0000", label="WARNING")
    doc.add_paragraph()

    # ── Section 1 ────────────────────────────────────────────────────────────
    add_heading(doc, "1. Project Identity", 1, (0x1A, 0x56, 0xAA))
    add_table(doc,
        ["Field", "Value"],
        [
            ["Name",             "Squelch"],
            ["CLI",              "python main.py  |  run_squelch.bat"],
            ["GitHub",           "github.com/dawardy/squelch"],
            ["License",          "GPL v3"],
            ["Language/Runtime", "Python 3.9+ / PyQt6"],
            ["Primary Platform", "Windows 10/11 x86"],
            ["Secondary",        "Linux (Ubuntu 22.04+, DragonOS)"],
            ["Hardware",         "ICOM IC-7100, RTL-SDR, RSP2Pro, USRP B200/B210, Hamlib 4.7.1"],
            ["Current Version",  "cat VERSION"],
        ],
        col_widths=[2.0, 4.5]
    )

    # ── Section 2 ────────────────────────────────────────────────────────────
    add_heading(doc, "2. Mission Statement", 1, (0x1A, 0x56, 0xAA))
    doc.add_paragraph(
        "Squelch is a unified amateur radio operations platform. It consolidates "
        "rig control, digital modes (FT8/FT4/WSPR/PSK31/RTTY), SDR waterfall, "
        "QSO logging, APRS, propagation analysis, and emergency communications "
        "tools into a single coherent application. The architecture is modular "
        "so first- and third-party add-on modules can extend capabilities "
        "without modifying core code."
    )

    # ── Section 3 ────────────────────────────────────────────────────────────
    add_heading(doc, "3. Architecture Summary", 1, (0x1A, 0x56, 0xAA))
    add_heading(doc, "3a. Layer Map", 2)
    add_table(doc,
        ["Layer", "Component", "Notes"],
        [
            ["Entry point",  "main.py",                        "setup_logging(force=True), _apply_theme_fixes()"],
            ["Config",       "core/config.py",                 "XDG config; %APPDATA%\\Squelch on Windows"],
            ["Safety FSM",   "core/safety.py",                 "AppState FSM; blocks unintended TX"],
            ["TX callsign",  "core/guest_op.py",               "operating_callsign(cfg) — FCC §97.119"],
            ["Themes",       "core/themes.py",                 "Dark/Light/High Contrast/Night; QSS tokens"],
            ["Terrain",      "core/terrain.py",                "SRTM; online (OpenTopoData) + offline (Amazon S3)"],
            ["Security",     "core/sanitize.py + credentials.py + netlog.py", "CSV injection, OS keyring, outbound log"],
            ["Plugins",      "core/plugins.py",                "Restricted to plugins/; MODULE_META contract"],
            ["UI base",      "ui/panel.py + panel_shell.py",   "SquelchPanel mixin; PanelShell workspace mode"],
            ["Main window",  "ui/main_window.py + 6 mixins",   "Profile, Network, GuestDemo, Firstrun, View, Location"],
            ["Tabs",         "ui/tabs/*.py",                   "All inherit SquelchPanel"],
            ["SDR",          "sdr/soapy_device.py + rtltcp",   "conda bridge; RTL-TCP fallback"],
            ["Network",      "network/pskreporter.py + repeaterbook.py", "PSKReporter, RepeaterBook, RadioID"],
        ],
        col_widths=[1.5, 2.2, 2.8]
    )

    add_heading(doc, "3b. AI Toolchain", 2)
    add_table(doc,
        ["Tool", "Role"],
        [
            ["Claude Code",     "Primary dev agent — architecture, refactoring, CI fixes"],
            ["GitHub Copilot",  "Inline autocomplete"],
            ["Perplexity",      "Research — protocols, algorithms, ham radio standards"],
            ["Zyte",            "Web scraping — callsign lookups, repeater data"],
            ["CodeScene",       "Code health scoring — target nothing below 7.0 actively changing"],
        ],
        col_widths=[2.0, 4.5]
    )

    # ── Section 4 ────────────────────────────────────────────────────────────
    add_heading(doc, "4. Feature Registry", 1, (0x1A, 0x56, 0xAA))
    add_heading(doc, "4a. MVP (Built)", 2)
    add_table(doc,
        ["ID", "Feature", "File", "Notes"],
        [
            ["RIG-01",  "IC-7100 CAT control",         "ui/tabs/rig_tab.py",               "Hamlib; VFO A/B, split"],
            ["DIG-01",  "FT8/FT4/WSPR decode list",    "ui/tabs/modes_tab.py",             "Via WSJT-X UDP"],
            ["LOG-01",  "QSO logging + ADIF export",   "ui/tabs/log_tab.py",               "SQLite backend"],
            ["MAP-01",  "APRS map + PSKReporter pins", "ui/tabs/map_tab.py",               "Leaflet HTML"],
            ["SDR-01",  "Waterfall + spectrum",         "ui/tabs/sdr_tab.py",               "SoapySDR / RTL-TCP"],
            ["PROP-01", "Band conditions + terrain",    "ui/tabs/band_conditions_tab.py",   "NOAA + SRTM"],
            ["SEC-01",  "Pentest suite",                "tests/test_security_pentest.py",   "Priya-38 / C-12"],
        ],
        col_widths=[0.8, 2.0, 2.5, 1.2]
    )
    add_heading(doc, "4b. Backlog (Priority Order)", 2)
    add_table(doc,
        ["ID", "Feature", "Status"],
        [
            ["C-01",       "SDR waterfall hardware test",            "Needs SoapySDR live hardware"],
            ["WORKSPACE-C","Snap zones",                             "Per-panel toolbars done; snap zones next"],
            ["THEME",      "Semantic-color refactor (~320 inline)",  "Dedicated sprint"],
            ["EIRP-2",     "Ionospheric path-loss model",            "Terrain done; model next"],
            ["MAP-2",      "More pin sources (Winlink, Local RF)",   "PSKReporter + FT8 done"],
            ["CS-SDR",     "sdr_tab.py CodeScene 5.68",             "_build_controls + _build_bottom_bar split next"],
        ],
        col_widths=[1.2, 3.0, 2.3]
    )

    # ── Section 5 ────────────────────────────────────────────────────────────
    add_heading(doc, "5. Plugin / Add-on Workflow", 1, (0x1A, 0x56, 0xAA))
    doc.add_paragraph("Every add-on module must satisfy this contract:")
    for item in [
        "All module code lives in plugins/<module_name>/",
        "__init__.py must expose MODULE_META dict and register(app) / unregister(app)",
        "register() returns a list of (panel_id, widget) tuples",
        "Same security rules as core — no shell=True, eval, exec, pickle",
        "All network calls use timeout= and core/netlog.record_connection()",
        "Must pass python qa_check.py before submission",
        "UI panels inherit SquelchPanel from ui/panel.py",
    ]:
        add_numbered(doc, item)

    add_code(doc,
        "MODULE_META = {\n"
        "    'name':        'My Module',\n"
        "    'version':     '0.1.0',\n"
        "    'author':      '[CALLSIGN]',\n"
        "    'description': 'One sentence.',\n"
        "    'squelch_min': '0.11.0',\n"
        "}\n\n"
        "def register(app) -> list:\n"
        "    # Return list of (panel_id, widget) tuples\n"
        "    return []\n\n"
        "def unregister(app) -> None:\n"
        "    pass  # stop threads, release hardware"
    )

    # ── Section 6 ────────────────────────────────────────────────────────────
    add_heading(doc, "6. DevSecOps Pipeline", 1, (0x1A, 0x56, 0xAA))
    add_heading(doc, "6a. Repository Structure", 2)
    for d in ["core/", "ui/", "sdr/", "network/", "modes/", "tests/", "plugins/", "docs/"]:
        add_bullet(doc, d)
    add_heading(doc, "6b. CI Gates", 2)
    add_table(doc,
        ["Gate", "Tool", "Pass Condition"],
        [
            ["__future__ order lint",  "test.yml custom check",      "No docstring before from __future__"],
            ["Undefined names",        "pyflakes",                   "python -m pyflakes ui core modes network exits 0"],
            ["Unit + integration",     "pytest",                     "QT_QPA_PLATFORM=offscreen pytest tests/ -q — 0 failures"],
            ["Method references",      "test_method_references.py",  "Every .connect(self._x) has _x defined"],
            ["Signal arity",           "test_signal_smoke.py",       "Slot arg count matches signal"],
            ["SAST security",          "test_security_pentest.py",   "23 tests pass; no shell=True/eval/pickle"],
            ["Syntax sweep",           "ast.parse in test.yml",      "All .py files parse clean"],
            ["CodeScene health",       "Manual sprint scan",         "Nothing below 7.0 actively changing"],
        ],
        col_widths=[1.8, 2.0, 2.7]
    )
    add_heading(doc, "6c. Sprint Cadence", 2)
    for item in [
        "Sprint 0: scaffolding only — 3 smoke tests pass before feature code",
        "Feature sprints: one logical area per sprint; stop after push, wait for feedback",
        "Packaging: python3 bump_version.py patch|minor|major → CHANGELOG entry → zip",
    ]:
        add_bullet(doc, item)

    # ── Section 7 ────────────────────────────────────────────────────────────
    add_heading(doc, "7. Stakeholder Framework", 1, (0x1A, 0x56, 0xAA))
    add_box(doc,
        "Stakeholder personas are INTERNAL. Do not commit to the public repository.",
        "FFFBE6", "CC8800", label="INTERNAL")
    add_table(doc,
        ["ID", "Persona", "Role", "Feedback Channel"],
        [
            ["C-02", "Dorothy-66",          "New ham, first-run setup",            "Sprint review, usability notes"],
            ["C-03", "Hank-72",             "Experienced operator, VFO/split",     "Sprint review"],
            ["C-04", "Marcus-34",           "Digital modes, FT8, local RF",        "Feature requests"],
            ["C-05", "Tyler-15",            "Visual/map focus",                    "UI feedback"],
            ["C-06", "Elena-50",            "Instructor, demo mode",               "Classroom scenario testing"],
            ["C-12", "Priya-38",            "Security analyst",                    "Pentest suite, C-12 compliance"],
            ["C-15", "Sam-19",              "Guest operator",                      "TX callsign / guest mode testing"],
            ["C-16", "RF Instructor",       "Spectrum training",                   "Signal ID, propagation demo"],
            ["C-19", "RF Security Analyst", "APRS anomaly detection",              "Protocol audit review"],
            ["C-20", "RF Researcher",       "RF recon, attack surface",            "Red-team review"],
            ["C-21", "Spectrum Trainer",    "Non-TX signal classification",        "Training scenario testing"],
            ["GOV",  "Government",          "Same standing as amateur stakeholders","Sprint reviews — always included"],
        ],
        col_widths=[0.6, 1.4, 2.2, 2.3]
    )

    # ── Section 8 ────────────────────────────────────────────────────────────
    add_heading(doc, "8. Session Handoff Template", 1, (0x1A, 0x56, 0xAA))
    doc.add_paragraph("Paste this block at the top of a new session or after a context reset:")
    add_code(doc,
        "## Handoff -- [DATE]\n"
        "**Last commit:** [git log --oneline -1]\n"
        "**Branch:** main\n"
        "**QA status:** [PASS / FAIL]\n"
        "**What was just completed:**\n"
        "  - [bullet]\n"
        "**What is in progress (do not restart):**\n"
        "  - [file, method, line range]\n"
        "**Next queued item:**\n"
        "  - [from Open backlog]\n"
        "**Known breakage:**\n"
        "  - [any failing tests]\n"
        "**Do not touch until resolved:**\n"
        "  - [file or feature]"
    )

    # ── Section 9 ────────────────────────────────────────────────────────────
    add_heading(doc, "9. CLAUDE.md Template", 1, (0x1A, 0x56, 0xAA))
    add_box(doc,
        "CLAUDE.md is gitignored and must never be committed to the public repository.",
        "FFE0E0", "CC0000", label="INTERNAL")
    doc.add_paragraph("Mandatory sections in every CLAUDE.md:")
    for s in [
        "What this is",
        "MANDATORY QA gate (qa_check.py)",
        "AI agent instructions (priority order, workflow rule, complexity cap)",
        "Architecture (file map)",
        "Design decisions — the WHY",
        "Security model (Priya-38 / C-12 compliance)",
        "Consumer personas",
        "Open backlog (priority order)",
        "DO NOT list",
        "CodeScene scores table",
        "AI toolchain",
        "Plugin / add-on module API",
        "Session handoff template",
        "Sprint 0 checklist",
        "Token efficiency rules",
        "CI gates",
    ]:
        add_bullet(doc, s)

    # ── Section 10 ───────────────────────────────────────────────────────────
    add_heading(doc, "10. Sprint 0 Checklist", 1, (0x1A, 0x56, 0xAA))
    for item in [
        "python qa_check.py exits 0 on main",
        "Branch created from clean main",
        "At least 3 smoke tests written and passing for target feature area",
        "Mock/stub layer in place so CI runs without physical hardware",
        "CLAUDE.md 'Last session summary' block updated",
        "CodeScene baseline scan run; scores table updated",
        "No functions > 60 lines in files scored below 7.0",
    ]:
        add_numbered(doc, item)

    # ── Section 11 ───────────────────────────────────────────────────────────
    add_heading(doc, "11. Token Efficiency Rules", 1, (0x1A, 0x56, 0xAA))
    for item in [
        "Read only files you need — use Grep/Glob before Read",
        "Never re-read a file just edited — edit tool confirms success",
        "One logical change per commit; no bundling of unrelated fixes",
        "Confirm qa_check.py passes before any commit",
        "Stop and wait for feedback after each push",
        "Perplexity: extract the one fact needed; do not paste walls of text into context",
        "When a method body exceeds ~60 lines, extract before continuing",
        "Do not chain multiple unrelated work batches autonomously",
    ]:
        add_bullet(doc, item)

    # ── Section 12 ───────────────────────────────────────────────────────────
    add_heading(doc, "12. Adapting This Template for Future Projects", 1, (0x1A, 0x56, 0xAA))
    add_box(doc,
        "This template was generated from Squelch's development history.",
        "E8F4E8", "2E7D32", label="NOTE")
    for item in [
        "Replace all Squelch-specific file paths, hardware targets, and persona IDs",
        "Keep Sections 8 (Handoff), 10 (Sprint 0), 11 (Token Efficiency), and 6b (CI gates) verbatim — these are universal scaffolding",
        "Update the CodeScene scores table after the first scan on the new codebase",
        "Add project-specific DO NOT rules as they are discovered in sprint retros",
        "The plugin/module API contract (Section 5) can be reused as-is for any PyQt6-based modular application",
    ]:
        add_numbered(doc, item)

    doc.save(OUT / "Squelch_Bootstrap_Template_INTERNAL.docx")
    print("Saved: Squelch_Bootstrap_Template_INTERNAL.docx")


# ── Document 2: Setup Walkthrough ────────────────────────────────────────────

def build_walkthrough():
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.25)
        section.right_margin  = Inches(1.25)

    title = doc.add_heading("Squelch — Developer Setup Walkthrough", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("From zero to first passing smoke test")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    sub.runs[0].font.size = Pt(12)
    doc.add_paragraph()

    # ── Section 1 ────────────────────────────────────────────────────────────
    add_heading(doc, "1. Prerequisites", 1, (0x1A, 0x56, 0xAA))
    add_table(doc,
        ["Item", "Where to Get It", "Notes"],
        [
            ["Python 3.12",       "python.org/downloads",                           "Check 'Add to PATH'; 3.12 recommended"],
            ["Git",               "git-scm.com",                                    "For repo clone and version tracking"],
            ["VS Code / PyCharm", "code.visualstudio.com",                          "GitHub Copilot works in both"],
            ["Hamlib 4.7.1",      "github.com/Hamlib/Hamlib/releases",              "Required for rig control"],
            ["SoapySDR (opt.)",   "PothosSDR — downloads.myriadrf.org/builds/PothosSDR/", "Only for SDR waterfall hardware test"],
        ],
        col_widths=[1.6, 2.8, 2.1]
    )

    # ── Section 2 ────────────────────────────────────────────────────────────
    add_heading(doc, "2. Repository Setup", 1, (0x1A, 0x56, 0xAA))
    steps = [
        ("Clone the repo:  ", "git clone https://github.com/dawardy/squelch.git"),
        ("Add AV exclusion FIRST", " — add the squelch/ folder to Windows Defender / Bitdefender / NETGEAR Armor exclusions before doing anything else"),
        ("Enter directory:  ", "cd squelch"),
        ("Run the installer:  ", "python installer.py"),
    ]
    for i, (bold, rest) in enumerate(steps, 1):
        p = doc.add_paragraph(style='List Number')
        r = p.add_run(bold); r.bold = True
        p.add_run(rest)
    doc.add_paragraph("The installer creates venv/, installs requirements.txt, checks external software, and creates run_squelch.bat.")

    # ── Section 3 ────────────────────────────────────────────────────────────
    add_heading(doc, "3. AI Toolchain Setup", 1, (0x1A, 0x56, 0xAA))

    tools = [
        ("Claude Code", "npm install -g @anthropic-ai/claude-code  (or pip install claude-code)\nRun: claude  in the project root.\nCLAUDE.md is gitignored and loaded automatically — it is your agent briefing."),
        ("GitHub Copilot", "Install the Copilot extension in VS Code.\nSign in with your GitHub account.\nWorks automatically alongside Claude Code."),
        ("Perplexity", "Web-based at perplexity.ai — no install needed.\nUse for research questions (protocols, algorithms, standards).\nPaste findings as comments or docs only — never raw web text into code."),
        ("Zyte", "pip install zyte-api\nStore API key as ZYTE_API_KEY via core/credentials.py → OS keyring.\nNever store in config.json."),
        ("CodeScene", "Web dashboard at codescene.io.\nConnect your GitHub repo.\nScan after each sprint; update the scores table in CLAUDE.md."),
    ]
    for name, detail in tools:
        add_heading(doc, f"3. {name}", 2)
        add_code(doc, detail)

    # ── Section 4 ────────────────────────────────────────────────────────────
    add_heading(doc, "4. Development Environment Verification", 1, (0x1A, 0x56, 0xAA))
    verif = [
        ("QA gate:  ", 'python qa_check.py  →  must print "QA PASSED — safe to package"'),
        ("Tests headless:  ", "QT_QPA_PLATFORM=offscreen pytest tests/ -q  (set env var on Windows)"),
        ("Security scan:  ", "pytest tests/test_security_pentest.py -v"),
        ("Undefined names:  ", "python -m pyflakes ui core modes network"),
        ("Launch app:  ", "python main.py  (or run_squelch.bat)"),
    ]
    for bold, rest in verif:
        p = doc.add_paragraph(style='List Number')
        r = p.add_run(bold); r.bold = True
        p.add_run(rest)

    # ── Section 5 ────────────────────────────────────────────────────────────
    add_heading(doc, "5. Hardware Setup", 1, (0x1A, 0x56, 0xAA))
    add_box(doc,
        "CI runs entirely without physical hardware. QT_QPA_PLATFORM=offscreen stubs out "
        "the display. SoapySDR is stubbed via sdr/soapy_device.py. You can develop and "
        "pass all tests without an IC-7100 or SDR dongle.",
        "E8F4FB", "1A6090", label="INFO")
    add_heading(doc, "IC-7100", 2)
    add_code(doc, "1. Install CP210x driver: silabs.com/developers/usb-to-uart-bridge-vcp-drivers\n2. Menu 066 (CI-V Baud Rate): 19200\n3. Menu 073 (CI-V USB Echo): OFF\n4. Rig tab → Model: ICOM IC-7100 → Connect")
    add_heading(doc, "RTL-SDR", 2)
    add_code(doc, "1. Zadig (zadig.akeo.ie): replace driver with WinUSB\n2. Install PothosSDR bundle\n3. pip install soapysdr")
    add_heading(doc, "SDRplay RSP series", 2)
    add_code(doc, "1. SDRplay API installer first (sdrplay.com/softwarehome)\n2. PothosSDR bundle\n3. pip install soapysdr")

    # ── Section 6 ────────────────────────────────────────────────────────────
    add_heading(doc, "6. First CI-Runnable Deliverable (Sprint 0)", 1, (0x1A, 0x56, 0xAA))
    add_box(doc,
        "Sprint 0 ends when three smoke tests pass in CI without physical hardware. "
        "Do not write feature code until this is true.",
        "FFFBE6", "CC8800", label="RULE")
    doc.add_paragraph("The three mandatory smoke tests:")

    add_heading(doc, "Test 1 — App imports without crash", 2)
    add_code(doc, "def test_main_imports():\n    import main\n    assert hasattr(main, 'main')")

    add_heading(doc, "Test 2 — QA gate passes", 2)
    add_code(doc, "def test_qa_check_passes():\n    import subprocess, sys\n    r = subprocess.run([sys.executable, 'qa_check.py'], capture_output=True)\n    assert r.returncode == 0")

    add_heading(doc, "Test 3 — Config loads", 2)
    add_code(doc, "def test_config_loads():\n    from core.config import SquelchConfig\n    cfg = SquelchConfig()\n    assert cfg is not None")

    doc.add_paragraph(
        "These three tests make CI green and give a runnable baseline before any "
        "feature work begins. The mock/stub layer means all three pass without "
        "an IC-7100 or SDR dongle connected."
    )

    # ── Section 7 ────────────────────────────────────────────────────────────
    add_heading(doc, "7. Writing Your First Add-On Module", 1, (0x1A, 0x56, 0xAA))
    for bold, rest in [
        ("Create:  ",     "plugins/my_module/__init__.py"),
        ("Add:  ",        "MODULE_META dict (name, version, author, description, squelch_min)"),
        ("Implement:  ",  "register(app) returning []  (empty for now)"),
        ("Implement:  ",  "unregister(app) as pass"),
        ("Verify:  ",     "python qa_check.py — must still exit 0"),
        ("Restart:  ",    "Module loads silently with no panels — expected at this stage"),
        ("Extend:  ",     "Add a SquelchPanel subclass and return it from register()"),
    ]:
        p = doc.add_paragraph(style='List Number')
        r = p.add_run(bold); r.bold = True
        p.add_run(rest)

    # ── Section 8 ────────────────────────────────────────────────────────────
    add_heading(doc, "8. Common First-Day Mistakes", 1, (0x1A, 0x56, 0xAA))
    mistakes = [
        ("Committed CLAUDE.md to the public repo",
         "It is gitignored for a reason. Personas and operational context stay internal."),
        ("Ran qa_check.py after the commit instead of before",
         "Always qa_check first. CI will catch it but a green local gate is faster."),
        ("Used QTimer.singleShot from a worker thread",
         "Silently never fires. Use pyqtSignal to cross thread boundaries."),
        ("Set 'from __future__ import annotations' after the module docstring",
         "The lint gate rejects this. It must be line 1."),
        ("Returned the inner widget from a QScrollArea",
         "The scroll area will be deleted. Return the scroll widget, not the inner container."),
    ]
    for title_text, detail in mistakes:
        add_box(doc, detail, "FFE0E0", "CC0000", label=f"MISTAKE: {title_text}")

    doc.save(OUT / "Squelch_Setup_Walkthrough.docx")
    print("Saved: Squelch_Setup_Walkthrough.docx")


if __name__ == "__main__":
    OUT.mkdir(parents=True, exist_ok=True)
    build_bootstrap()
    build_walkthrough()
    print("Done.")
